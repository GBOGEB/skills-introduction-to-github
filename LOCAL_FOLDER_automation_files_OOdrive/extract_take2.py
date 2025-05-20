import re
import json
from docx import Document
import yaml  # Ensure PyYAML is installed

# === File Paths ===
DOCX_TEMPLATE_PATH = r'C:\Users\gbonthuy\Downloads\MASTER_1305_0951.docx'
OUTPUT_JSON_PATH = r'C:\Users\gbonthuy\Downloads\qqq_extracted.json'
OUTPUT_YAML_PATH = r'C:\Users\gbonthuy\Downloads\qqq_extracted.yaml'
OUTPUT_MD_PATH   = r'C:\Users\gbonthuy\Downloads\qqq_extracted.md'
OUTPUT_PY_PATH   = r'C:\Users\gbonthuy\Downloads\qqq_extracted.py'
OUTLINE_JSON_PATH = r'C:\Users\gbonthuy\Downloads\MASTER_numbered_outline.json'

# === Load outline map (paragraph indices + titles) ===
with open(OUTLINE_JSON_PATH, encoding="utf-8") as f:
    outline = json.load(f)

# === Build paragraph-index-to-outline map ===
outline_map = {}
for entry in outline:
    outline_map[entry["paragraph_index"]] = {
        "outline_number": entry.get("number"),
        "outline_title": entry.get("title")
    }

# === QQQ regex ===
qqq_re = re.compile(r"QQQ[_\s]*(\d{3,5})\s+(.*)", re.IGNORECASE)

# === State tracking ===
results = []
last_outline = None

# === Load .docx ===
doc = Document(DOCX_TEMPLATE_PATH)

# === Scan paragraphs with position ===
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    # Track last valid outline using paragraph indices
    if idx in outline_map:
        last_outline = outline_map[idx]

    # Debug: Print each paragraph for insight
    print(f"Paragraph {idx}: {text}")

    # Match QQQ_### requirements
    if "QQQ" in text and any(char.isdigit() for char in text):
        parts = text.split("QQQ_")
        if len(parts) > 1:
            num_part = parts[1].split()[0].strip(":.-)–")
            try:
                qnum = int(num_part)
                body = text.split(num_part, 1)[-1].strip(":.-)– ")
                results.append({
                    "id": f"QQQ_{qnum:03d}",
                    "number": qnum,
                    "text": body,
                    "outline_number": last_outline.get("outline_number") if last_outline else None,
                    "outline_title": last_outline.get("outline_title") if last_outline else None
                })
            except ValueError:
                pass  # Skip misformatted lines
    else:
        print(f"Paragraph {idx}: No match for QQQ regex")

# Debug: Print paragraphs containing "QQQ"
for i, para in enumerate(doc.paragraphs):
    if "QQQ" in para.text:
        print(i, para.text.strip())

# Debug: Print style and text of paragraphs containing "QQQ"
for para in doc.paragraphs:
    if "QQQ" in para.text:
        print(para.style.name, para.text.strip())

# === Phase 2: Exporting Outputs ===

# Output 1: JSON
with open(OUTPUT_JSON_PATH, "w", encoding="utf-8") as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"✅ Extracted {len(results)} QQQ requirements with outline context into {OUTPUT_JSON_PATH}")

# Output 2: YAML
with open(OUTPUT_YAML_PATH, "w", encoding="utf-8") as f:
    yaml.dump(results, f, sort_keys=False)
print(f"✅ YAML export completed: {OUTPUT_YAML_PATH}")

# Output 3: Markdown Table
with open(OUTPUT_MD_PATH, "w", encoding="utf-8") as f:
    f.write("| ID | Text (Preview) | Outline |\n")
    f.write("|----|----------------|---------|\n")
    for r in results:
        preview = (r['text'][:80] + '...') if len(r['text']) > 80 else r['text']
        outline_info = f"{r['outline_number'] or ''} {r['outline_title'] or ''}".strip()
        f.write(f"| {r['id']} | {preview} | {outline_info} |\n")
print(f"✅ Markdown export completed: {OUTPUT_MD_PATH}")

# Output 4: Python List
with open(OUTPUT_PY_PATH, "w", encoding="utf-8") as f:
    f.write("qqq_requirements = [\n")
    for r in results:
        outline_info = f"{r['outline_number'] or ''} {r['outline_title'] or ''}".strip()
        f.write(f"  {{'id': '{r['id']}', 'number': {r['number']}, 'text': '''{r['text']}''', 'outline': '{outline_info}'}},\n")
    f.write("]\n")
print(f"✅ Python list export completed: {OUTPUT_PY_PATH}")
